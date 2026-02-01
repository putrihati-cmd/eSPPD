"""
LPJ Document Generator - Formatted Table-Based Report
Generates LPJ (Laporan Perjalanan Dinas) with specific table formatting.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional
import uuid
import subprocess


class LPJDocumentGenerator:
    """
    Generate LPJ document with proper Indonesian government format.
    Follows standard table-based layout as specified in lpj.md.
    """
    
    def __init__(self, output_dir: str = "./generated"):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)
    
    def generate(self, data: Dict) -> str:
        """
        Generate LPJ document in Word format (.docx)
        
        Args:
            data: Dictionary containing all required fields
            
        Returns:
            Path to generated file
        """
        doc = Document()
        
        # Setup margins
        sections = doc.sections[0]
        sections.top_margin = Cm(2)
        sections.bottom_margin = Cm(2)
        sections.left_margin = Cm(2.5)
        sections.right_margin = Cm(2)
        
        # Header (Title)
        self._add_title(doc, "LAPORAN PERJALANAN DINAS")
        
        # Main table
        self._create_main_table(doc, data)
        
        # Add spacing
        doc.add_paragraph()
        
        # Signature section
        self._create_signature_table(doc, data)
        
        # Generate filename
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        unique_id = str(uuid.uuid4())[:8]
        nip_safe = data.get('nip', 'unknown').replace('/', '_').replace(' ', '')
        filename = f"LPJ_{nip_safe}_{timestamp}_{unique_id}.docx"
        output_path = self.output_dir / filename
        
        doc.save(output_path)
        
        return str(output_path)
    
    def _add_title(self, doc, title: str):
        """Add centered bold title"""
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(title)
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.name = 'Times New Roman'
    
    def _create_main_table(self, doc, data: Dict):
        """Create the main LPJ table with proper formatting"""
        table = doc.add_table(rows=9, cols=3)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Set column widths
        widths = [Cm(0.8), Cm(5), Cm(10)]
        for i, width in enumerate(widths):
            for cell in table.columns[i].cells:
                cell.width = width
        
        # Row 1: Employee info
        self._set_cell(table, 0, 0, "1", bold=True)
        self._set_cell(table, 0, 1, "Pegawai yang melaksanakan SPD\nNIP")
        self._set_cell(table, 0, 2, f"{data.get('name', '')}\n{data.get('nip', '')}")
        
        # Row 2: Rank
        self._set_cell(table, 1, 0, "")
        self._set_cell(table, 1, 1, "a) Pangkat/Golongan")
        rank = data.get('rank', '')
        gol = data.get('golongan', data.get('class', ''))
        self._set_cell(table, 1, 2, f"{rank} ({gol})" if rank else gol)
        
        # Row 3: Position
        self._set_cell(table, 2, 0, "")
        self._set_cell(table, 2, 1, "b) Jabatan")
        self._set_cell(table, 2, 2, data.get('position', data.get('jabatan', '')))
        
        # Row 4: Destination
        self._set_cell(table, 3, 0, "2", bold=True)
        self._set_cell(table, 3, 1, "Tempat tujuan")
        self._set_cell(table, 3, 2, data.get('destination', data.get('tujuan', '')))
        
        # Row 5: Purpose
        self._set_cell(table, 4, 0, "3", bold=True)
        self._set_cell(table, 4, 1, "Maksud perjalanan dinas")
        self._set_cell(table, 4, 2, data.get('purpose', data.get('keperluan', '')))
        
        # Row 6: Date
        self._set_cell(table, 5, 0, "4", bold=True)
        self._set_cell(table, 5, 1, "Hari / tanggal")
        day = data.get('day', '')
        date_range = data.get('date_range', f"{data.get('tanggal_berangkat', '')} s.d. {data.get('tanggal_kembali', '')}")
        self._set_cell(table, 5, 2, f"{day}, {date_range}" if day else date_range)
        
        # Row 7: Duration
        self._set_cell(table, 6, 0, "5", bold=True)
        self._set_cell(table, 6, 1, "Lama perjalanan Dinas")
        duration = data.get('duration', data.get('lama_perjalanan', 1))
        self._set_cell(table, 6, 2, f"{duration} Hari")
        
        # Row 8: Basis (with sub-table)
        self._set_cell(table, 7, 0, "6", bold=True)
        self._set_cell(table, 7, 1, "Dasar perjalanan dinas")
        
        # Create sub-content for dasar perjalanan
        basis_content = f"""a) Undangan: {data.get('invitation', '-')}
b) Surat Tugas: {data.get('assignment_letter', data.get('nomor_surat_tugas', ''))}
c) SPD: {data.get('sppd_number', data.get('nomor_sppd', ''))}"""
        self._set_cell(table, 7, 2, basis_content)
        
        # Row 9: Report content
        self._set_cell(table, 8, 0, "7", bold=True)
        self._set_cell(table, 8, 1, "Isi laporan perjalanan Dinas")
        
        # Build report content
        outputs = data.get('outputs', [])
        if isinstance(outputs, str):
            outputs = [outputs]
        
        outputs_text = "\n".join([f"   {i+1}. {out}" for i, out in enumerate(outputs)]) if outputs else "-"
        
        report_content = f"""a) Tanggal berangkat: {data.get('departure_date', data.get('tanggal_berangkat', ''))}

b) Isi perjalanan dinas:
{data.get('report_content', data.get('kegiatan', ''))}

c) Output perjalanan dinas:
{outputs_text}

d) Tanggal tiba di tempat kedudukan: {data.get('return_date', data.get('tanggal_kembali', ''))}"""
        
        self._set_cell(table, 8, 2, report_content)
    
    def _create_signature_table(self, doc, data: Dict):
        """Create signature section at bottom"""
        sig_table = doc.add_table(rows=5, cols=2)
        sig_table.style = 'Table Grid'
        sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Headers
        self._set_cell(sig_table, 0, 0, "Mengetahui\nAtasan Langsung", center=True)
        
        report_place = data.get('report_place', 'Purwokerto')
        report_date = data.get('report_date', datetime.now().strftime('%d %B %Y'))
        self._set_cell(sig_table, 0, 1, f"{report_place}, {report_date}\nYang melaporkan", center=True)
        
        # Empty rows for signature space
        for i in range(1, 4):
            self._set_cell(sig_table, i, 0, "")
            self._set_cell(sig_table, i, 1, "")
        
        # Names
        self._set_cell(sig_table, 4, 0, f"\n{data.get('superior_name', data.get('atasan_nama', ''))}\nNIP. {data.get('superior_nip', data.get('atasan_nip', ''))}", center=True, bold=True)
        self._set_cell(sig_table, 4, 1, f"\n{data.get('employee_name', data.get('name', ''))}\nNIP. {data.get('nip', '')}", center=True, bold=True)
    
    def _set_cell(self, table, row: int, col: int, text: str, bold: bool = False, center: bool = False):
        """Set cell content with formatting"""
        cell = table.rows[row].cells[col]
        cell.text = text
        
        for paragraph in cell.paragraphs:
            if center:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)
                if bold:
                    run.font.bold = True
    
    def convert_to_pdf(self, docx_path: str) -> str:
        """
        Convert DOCX to PDF using LibreOffice
        
        Args:
            docx_path: Path to DOCX file
            
        Returns:
            Path to generated PDF
        """
        import os
        docx_path = Path(docx_path)
        pdf_path = docx_path.with_suffix('.pdf')
        
        # Common LibreOffice paths on Windows
        soffice_paths = [
            'soffice', # If in PATH
            r'C:\Program Files\LibreOffice\program\soffice.exe',
            r'C:\Program Files (x86)\LibreOffice\program\soffice.exe',
            r'C:\laragon\bin\libreoffice\program\soffice.exe',
        ]
        
        soffice_bin = None
        for path in soffice_paths:
            if path != 'soffice' and os.path.exists(path):
                soffice_bin = path
                break
        
        if not soffice_bin:
            soffice_bin = 'soffice'

        try:
            subprocess.run([
                soffice_bin,
                '--headless',
                '--convert-to', 'pdf',
                '--outdir', str(docx_path.parent),
                str(docx_path)
            ], check=True, capture_output=True)
            
            return str(pdf_path)
        except subprocess.CalledProcessError as e:
            raise RuntimeError(f"PDF conversion failed: {e.stderr.decode()}")
        except FileNotFoundError:
            raise RuntimeError("LibreOffice not found in common paths. Please install LibreOffice and add 'soffice' to PATH.")
